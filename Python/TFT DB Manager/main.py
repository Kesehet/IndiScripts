# import the required libraries
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import pickle
from googleapiclient import errors
import os.path
import base64
import email
from bs4 import BeautifulSoup
from docx import Document
import pandas as pd
import datetime
import glob
import os
# Define the SCOPES. If modifying it, delete the token.pickle file.
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']

def getEmails():
  # Variable creds will store the user access token.
  # If no valid token found, we will create one.
  creds = None

  # The file token.pickle contains the user access token.
  # Check if it exists
  if os.path.exists('token.pickle'):

    # Read the token from the file and store it in the variable creds
    with open('token.pickle', 'rb') as token:
      creds = pickle.load(token)

  # If credentials are not available or are invalid, ask the user to log in.
  if not creds or not creds.valid:
    if creds and creds.expired and creds.refresh_token:
      creds.refresh(Request())
    else:
      flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
      creds = flow.run_local_server(port=0)

    # Save the access token in token.pickle file for the next run
    with open('token.pickle', 'wb') as token:
      pickle.dump(creds, token)

  # Connect to the Gmail API
  service = build('gmail', 'v1', credentials=creds)

  # request a list of all the messages
  result = service.users().messages().list(maxResults=20,userId='me').execute()

  # We can also pass maxResults to get any number of emails. Like this:
  # result = service.users().messages().list(maxResults=200, userId='me').execute()
  messages = result.get('messages')

  # messages is a list of dictionaries where each dictionary contains a message id.

  # iterate through all the messages
  for msg in messages:
    # Get the message from its id
    txt = service.users().messages().get(userId='me', id=msg['id']).execute()
    rytperson = False
    try:
        payload = txt['payload']
        headers = payload['headers']

          # Look for Subject and Sender Email in the headers
        for d in headers:
            print(d)
            if d['name'] == 'From':
                sender = d['value']
                if sender.find("newperspectivemedia") >= 0:
                    rytperson = True
            print("\n")
    except:
        pass
    # Use try-except to avoid any Errors
    try:
      # Get value of 'payload' from dictionary 'txt'
      payload = txt['payload']
      headers = payload['headers']

      # Look for Subject and Sender Email in the headers
      for d in headers:
                            print(d)
                            if d['name'] == 'Subject':
                                        subject = d['value']
                            if d['name'] == 'From':
                                        sender = d['value']
                            print("\n")

      # The Body of the message is in Encrypted format. So, we have to decode it.
      # Get the data and decode it with base 64 decoder.
      parts = payload.get('parts')[0]
      data = parts['body']['data']
      data = data.replace("-","+").replace("_","/")
      decoded_data = base64.b64decode(data)
      

      # Now, the data obtained is in lxml. So, we will parse
      # it with BeautifulSoup library
      soup = BeautifulSoup(decoded_data , "lxml")
      
      body = soup.body()
    
      # Printing the subject, sender's email and message
      #print("Subject: ", subject)
      #print("From: ", sender)
      #print("Message: ", body)
      print('\n')
    except:
      pass
    try:
        parts = [txt['payload']]
        while parts:
            part = parts.pop()
            if part.get('parts'):
                parts.extend(part['parts'])
            if part.get('filename'):
                if 'data' in part['body']:
                    file_data = base64.urlsafe_b64decode(part['body']['data'].encode('UTF-8'))
                    #self.stdout.write('FileData for %s, %s found! size: %s' % (message['id'], part['filename'], part['size']))
                elif 'attachmentId' in part['body']:
                    attachment = service.users().messages().attachments().get(
                        userId="me", messageId=msg['id'], id=part['body']['attachmentId']
                    ).execute()
                    file_data = base64.urlsafe_b64decode(attachment['data'].encode('UTF-8'))
                    #self.stdout.write('FileData for %s, %s found! size: %s' % (message['id'], part['filename'], attachment['size']))
                else:
                    file_data = None
                if file_data:
                    store_dir = "store_dir/"
                    path = ''.join([store_dir, part['filename']])
                    
                    with open(path, 'wb') as f:
                        f.write(file_data)
    except errors.HttpError as error:
        print ('An error occurred:'+error)
    if rytperson:
        #break
      pass

def createCSVfromWord():
    dir_name = "store_dir/"
    test = os.listdir(dir_name)
    docname = ""
    for item in test:
        if item.endswith(".docx"):
            docname = item
            break
    doc = Document('store_dir/'+str(docname))
      
    # print the list of paragraphs in the document
    shift = "Morning" if datetime.datetime.now().hour <= 12 else "Evening"
    stories = []
    stories.append({"snoPerDay":"______","Headline":"_______","time":str(datetime.datetime.now().day)+"-"+str(datetime.datetime.now().month)+"-"+str(datetime.datetime.now().year)+" - "+shift})
    for i in range(len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if i == 0:
            print (t)
            stories.append({"snoPerDay":len(stories),"Headline":t,"time":shift})
        if doc.paragraphs[i].text == "" and doc.paragraphs[i+1].text == "" and doc.paragraphs[i+2].text != "":
            print (doc.paragraphs[i+2].text)
            stories.append({"snoPerDay":len(stories),"Headline":doc.paragraphs[i+2].text.strip(),"time":shift})
    print(stories)
    df = pd.DataFrame(stories)
    df.to_csv("new_list.csv",index=False)

    for item in test:
        if item.endswith(".docx") and item.find("TFT") == -1:
            os.remove(os.path.join(dir_name, item))

def merger():
    df = pd.concat(map(pd.read_csv, glob.glob(os.path.join('', "*.csv") ) ) ).drop_duplicates().reset_index(drop=True)
    df.to_csv("main_list.csv",index=False)


getEmails()
createCSVfromWord()
merger()


